from __future__ import annotations

import json
import os
import sys
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Tuple
from urllib.parse import urlparse
import re

import requests
from docx import Document
from dotenv import load_dotenv

# ==========================
# USER EDITABLE CONFIG
# ==========================
OUTPUT_DIR = "downloads"
LINKS = [
    "https://ulii.org/akn/ug/act/2015/3/eng%402015-03-06.pdf",
    "https://bills.parliament.ug/attachments/Laws%20of%20Uganda%20%28Acts%29%20-%20THE%20FINANCE%20ACT%202013.pdf",
    "https://ulii.org/akn/ug/act/statute/1995/constitution/eng.pdf",
    # "https://example.com/file1.pdf",
]
# ==========================

KEYWORDS = ["woman", "gender"]
REPORT_FILE = "keyword_paragraph_report.docx"
OCR_TIMEOUT_SECONDS = 120
POLL_INTERVAL_SECONDS = 2


@dataclass
class DownloadResult:
    downloaded: int = 0
    skipped: int = 0
    failed: int = 0
    files: List[Path] = None

    def __post_init__(self) -> None:
        if self.files is None:
            self.files = []


def derive_filename(url: str) -> str:
    parsed = urlparse(url)
    name = Path(parsed.path).name
    if not name:
        name = f"download_{int(time.time())}.pdf"
    if not name.lower().endswith(".pdf"):
        name = f"{name}.pdf"
    return name


def ensure_output_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def stream_download(url: str, dest: Path) -> bool:
    try:
        with requests.get(url, stream=True, timeout=30) as response:
            if response.status_code != 200:
                print(f"[FAIL] {url} -> HTTP {response.status_code}")
                return False

            total_size = response.headers.get("Content-Length")
            size_label = f"{int(total_size) / (1024 * 1024):.2f} MB" if total_size else "unknown size"
            print(f"[GET ] {dest.name} ({size_label})")

            with dest.open("wb") as file_handle:
                downloaded_bytes = 0
                for chunk in response.iter_content(chunk_size=64 * 1024):
                    if not chunk:
                        continue
                    file_handle.write(chunk)
                    downloaded_bytes += len(chunk)
                    if total_size:
                        percent = (downloaded_bytes / int(total_size)) * 100
                        print(f"       {dest.name}: {percent:6.2f}%", end="\r")
            if total_size:
                print(" " * 80, end="\r")
            print(f"[ OK ] Downloaded {dest.name}")
            return True
    except requests.RequestException as exc:
        print(f"[FAIL] {url} -> {exc}")
        return False


def download_pdfs(output_dir: Path, links: List[str]) -> DownloadResult:
    result = DownloadResult()
    ensure_output_dir(output_dir)

    for url in links:
        filename = derive_filename(url)
        destination = output_dir / filename

        if destination.exists():
            print(f"[SKIP] {filename} already exists")
            result.skipped += 1
            result.files.append(destination)
            continue

        success = stream_download(url, destination)
        if success:
            result.downloaded += 1
            result.files.append(destination)
        else:
            result.failed += 1

    return result


def get_azure_settings() -> Tuple[str, str]:
    endpoint = os.getenv("AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT", "").strip()
    key = os.getenv("AZURE_DOCUMENT_INTELLIGENCE_KEY", "").strip()
    if not endpoint or not key:
        raise RuntimeError(
            "Set AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT and AZURE_DOCUMENT_INTELLIGENCE_KEY environment variables."
        )
    return endpoint.rstrip("/"), key


def analyze_pdf_paragraphs(pdf_path: Path, endpoint: str, key: str) -> List[str]:
    analyze_urls = [
        f"{endpoint}/documentintelligence/documentModels/prebuilt-read:analyze?api-version=2024-02-29-preview",
        f"{endpoint}/formrecognizer/documentModels/prebuilt-read:analyze?api-version=2023-07-31",
    ]
    headers = {
        "Ocp-Apim-Subscription-Key": key,
        "Content-Type": "application/octet-stream",
    }
    start_response = None
    last_error: Exception | None = None
    for analyze_url in analyze_urls:
        try:
            with pdf_path.open("rb") as file_handle:
                candidate_response = requests.post(
                    analyze_url, headers=headers, data=file_handle, timeout=60
                )
            if candidate_response.status_code in (404, 405):
                continue
            candidate_response.raise_for_status()
            start_response = candidate_response
            break
        except requests.RequestException as exc:
            last_error = exc
            continue

    if start_response is None:
        if last_error:
            raise last_error
        raise RuntimeError("No compatible analyze endpoint found for this resource.")

    operation_url = start_response.headers.get("Operation-Location")
    if not operation_url:
        raise RuntimeError("Azure response missing Operation-Location header.")

    poll_headers = {"Ocp-Apim-Subscription-Key": key}
    started_at = time.time()

    while True:
        poll_response = requests.get(operation_url, headers=poll_headers, timeout=30)
        poll_response.raise_for_status()
        payload = poll_response.json()
        status = payload.get("status", "").lower()

        if status == "succeeded":
            analyze_result = payload.get("analyzeResult", {})
            paragraphs = analyze_result.get("paragraphs", [])
            paragraph_texts = [str(item.get("content", "")).strip() for item in paragraphs if item.get("content")]
            if paragraph_texts:
                return paragraph_texts

            pages = analyze_result.get("pages", [])
            lines_fallback: List[str] = []
            for page in pages:
                lines = page.get("lines", [])
                for line in lines:
                    content = str(line.get("content", "")).strip()
                    if content:
                        lines_fallback.append(content)
            return lines_fallback

        if status == "failed":
            raise RuntimeError(f"OCR failed for {pdf_path.name}: {json.dumps(payload)}")

        if time.time() - started_at > OCR_TIMEOUT_SECONDS:
            raise TimeoutError(f"OCR timed out for {pdf_path.name}")

        time.sleep(POLL_INTERVAL_SECONDS)


def build_keyword_index(paragraphs_by_file: Dict[str, List[str]], keywords: List[str]) -> Dict[str, List[Dict[str, str]]]:
    index: Dict[str, List[Dict[str, str]]] = {keyword: [] for keyword in keywords}
    lowered_keywords = [keyword.lower() for keyword in keywords]

    for file_name, paragraphs in paragraphs_by_file.items():
        for paragraph in paragraphs:
            paragraph_lower = paragraph.lower()
            for original_keyword, lowered_keyword in zip(keywords, lowered_keywords):
                if lowered_keyword in paragraph_lower:
                    index[original_keyword].append({"file": file_name, "paragraph": paragraph})
    return index


def clean_paragraph(text: str) -> str:
    # Remove OCR replacement artifacts and collapse whitespace for report readability.
    cleaned = text.replace("\ufffd", "").replace("□", "")
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned


def write_word_report(report_path: Path, keyword_index: Dict[str, List[Dict[str, str]]]) -> None:
    document = Document()
    document.add_heading("Keyword Paragraph Report", level=0)
    document.add_paragraph(f"Generated at: {time.strftime('%Y-%m-%d %H:%M:%S')}")

    for keyword, matches in keyword_index.items():
        document.add_heading(f"Keyword: {keyword}", level=1)
        document.add_paragraph(f"Matches found: {len(matches)}")
        if not matches:
            document.add_paragraph("No matching paragraphs found.")
            continue
        for match in matches:
            paragraph_text = clean_paragraph(match["paragraph"])
            if not paragraph_text:
                continue
            source_line = document.add_paragraph()
            source_line.add_run(match["file"]).bold = True
            document.add_paragraph(paragraph_text)
            document.add_paragraph("")

    document.save(str(report_path))


def run_keyword_ocr(
    links: List[str],
    keywords: List[str],
    output_dir: Path,
    endpoint: str,
    key: str,
    uploaded_files: List[Path] = None,
) -> Dict[str, any]:
    """
    Library function for keyword OCR processing.
    Returns structured data instead of printing and exiting.
    """
    result = {
        "download_summary": {"downloaded": 0, "skipped": 0, "failed": 0},
        "ocr_summary": {"processed": 0, "failed": 0, "errors": []},
        "keyword_index": {},
        "report_path": None,
    }

    ensure_output_dir(output_dir)
    files_to_process: List[Path] = []

    if links:
        download_result = download_pdfs(output_dir, links)
        result["download_summary"] = {
            "downloaded": download_result.downloaded,
            "skipped": download_result.skipped,
            "failed": download_result.failed,
        }
        files_to_process.extend([f for f in download_result.files if f.exists()])

    if uploaded_files:
        files_to_process.extend(uploaded_files)

    if not files_to_process:
        return result

    paragraphs_by_file: Dict[str, List[str]] = {}
    for pdf_file in files_to_process:
        try:
            paragraphs_by_file[pdf_file.name] = analyze_pdf_paragraphs(pdf_file, endpoint, key)
            result["ocr_summary"]["processed"] += 1
        except Exception as exc:
            result["ocr_summary"]["failed"] += 1
            result["ocr_summary"]["errors"].append({"file": pdf_file.name, "error": str(exc)})

    if paragraphs_by_file:
        keyword_index = build_keyword_index(paragraphs_by_file, keywords)
        result["keyword_index"] = {
            kw: [{"file": m["file"], "paragraph": clean_paragraph(m["paragraph"])} for m in matches]
            for kw, matches in keyword_index.items()
        }
        report_path = output_dir / REPORT_FILE
        write_word_report(report_path, keyword_index)
        result["report_path"] = str(report_path)

    return result


def main() -> int:
    load_dotenv(Path(__file__).resolve().parent / ".env.local")

    if not LINKS:
        print("No LINKS configured. Add direct PDF URLs at the top of the script.")
        return 1

    output_dir = Path(OUTPUT_DIR)

    try:
        endpoint, key = get_azure_settings()
    except RuntimeError as exc:
        print(f"[ERROR] {exc}")
        return 1

    result = run_keyword_ocr(LINKS, KEYWORDS, output_dir, endpoint, key)

    print(
        f"\nDownload summary -> downloaded: {result['download_summary']['downloaded']}, "
        f"skipped: {result['download_summary']['skipped']}, failed: {result['download_summary']['failed']}"
    )
    print(
        f"\nOCR summary -> processed: {result['ocr_summary']['processed']}, "
        f"failed: {result['ocr_summary']['failed']}"
    )

    if result["ocr_summary"]["errors"]:
        print("\nOCR Errors:")
        for error in result["ocr_summary"]["errors"]:
            print(f"  - {error['file']}: {error['error']}")

    if result["keyword_index"]:
        print("\nKeyword index summary:")
        for keyword, matches in result["keyword_index"].items():
            print(f"  - {keyword}: {len(matches)} matches")
        print(f"\nWord report created: {result['report_path']}")
        return 0
    else:
        print("No OCR output generated.")
        return 1


if __name__ == "__main__":
    sys.exit(main())
